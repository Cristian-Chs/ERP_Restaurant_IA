"""
Script para generar documentación completa del sistema en formato Word (.docx)
Ejecutar: python generate_documentation.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os

def create_documentation():
    """Genera el documento Word completo del sistema"""
    
    doc = Document()
    
    # === PORTADA ===
    title = doc.add_heading('Sistema de Restaurante 4 Sabores', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Documentación Técnica Completa')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    version_info = doc.add_paragraph(f'Versión: 2.0\nFecha: {datetime.now().strftime("%d de %B de %Y")}')
    version_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # === TABLA DE CONTENIDOS ===
    doc.add_heading('Tabla de Contenidos', 1)
    toc_items = [
        '1. Introducción al Sistema',
        '2. Arquitectura del Sistema',
        '3. Modelos de Base de Datos',
        '4. Módulos del Backend',
        '5. API Reference Completo',
        '6. Automatizaciones con Pandas',
        '7. Frontend - Componentes y Páginas',
        '8. Bot de Telegram',
        '9. Manual de Usuario',
        '10. Configuración y Deployment'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # === 1. INTRODUCCIÓN ===
    doc.add_heading('1. Introducción al Sistema', 1)
    
    doc.add_paragraph(
        'El Sistema de Restaurante 4 Sabores es una solución integral diseñada para modernizar '
        'las operaciones del restaurante mediante la automatización de procesos clave. Integra '
        'un Bot de Telegram para pedidos automatizados con una Aplicación Web React de alto '
        'rendimiento para la gestión de cocina y administración.'
    )
    
    doc.add_heading('Objetivos del Proyecto', 2)
    objectives = [
        'Automatizar el proceso de pedidos mediante Telegram y Web',
        'Optimizar la gestión de inventario con recálculo dinámico de costos',
        'Proporcionar análisis inteligente de negocio con Pandas',
        'Facilitar la toma de decisiones mediante reportes avanzados',
        'Mejorar la experiencia del cliente con recomendaciones personalizadas'
    ]
    
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_heading('Stack Tecnológico', 2)
    
    # Backend
    doc.add_heading('Backend', 3)
    backend_tech = [
        'Framework: Django 5.2.8',
        'API: Django Rest Framework (DRF)',
        'Base de Datos: SQLite (Desarrollo) / PostgreSQL (Producción)',
        'Analytics: Pandas 2.3.3, NumPy 2.3.5',
        'IA/ML: Scikit-learn, Sentence Transformers',
        'Autenticación: JWT (Djoser) + Telegram Auth',
        'Reportes: ReportLab (PDF), OpenPyXL (Excel), python-docx (Word)'
    ]
    
    for tech in backend_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    # Frontend
    doc.add_heading('Frontend', 3)
    frontend_tech = [
        'Framework: React 18 con Vite',
        'Estilos: CSS Modules',
        'Gestión de Estado: React Hooks (Context API)',
        'Rutas: React Router DOM 6',
        'HTTP Client: Fetch API'
    ]
    
    for tech in frontend_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_page_break()
    
    # === 2. ARQUITECTURA ===
    doc.add_heading('2. Arquitectura del Sistema', 1)
    
    doc.add_paragraph(
        'El sistema sigue una arquitectura Cliente-Servidor con separación clara entre '
        'frontend (React) y backend (Django). La comunicación se realiza mediante API REST.'
    )
    
    doc.add_heading('Componentes Principales', 2)
    
    components = {
        'Frontend React': 'Interfaz de usuario para clientes, administradores y cocina',
        'Backend Django': 'API REST, lógica de negocio y gestión de datos',
        'Bot de Telegram': 'Canal alternativo para pedidos y notificaciones',
        'Base de Datos': 'Almacenamiento persistente de productos, pedidos y usuarios',
        'Motor de Analytics': 'Procesamiento de datos con Pandas para insights de negocio',
        'Motor de IA': 'Recomendaciones personalizadas basadas en ML'
    }
    
    for component, description in components.items():
        p = doc.add_paragraph()
        p.add_run(f'{component}: ').bold = True
        p.add_run(description)
    
    doc.add_page_break()
    
    # === 3. MODELOS DE BASE DE DATOS ===
    doc.add_heading('3. Modelos de Base de Datos', 1)
    
    doc.add_heading('App: core', 2)
    
    # Ingredient
    doc.add_heading('Modelo: Ingredient (Ingrediente)', 3)
    doc.add_paragraph(
        'Representa las materias primas utilizadas en la preparación de productos. '
        'Incluye información de costos para cálculos dinámicos.'
    )
    
    ingredient_fields = [
        ('nombre', 'CharField(100)', 'Nombre del ingrediente (único)'),
        ('disponible_como_extra', 'BooleanField', 'Si puede agregarse como extra'),
        ('cost', 'DecimalField', 'Costo por unidad'),
        ('unit', 'CharField(5)', 'Unidad de medida (kg, l, und)')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Campo'
    hdr_cells[1].text = 'Tipo'
    hdr_cells[2].text = 'Descripción'
    
    for field, ftype, desc in ingredient_fields:
        row_cells = table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = ftype
        row_cells[2].text = desc
    
    doc.add_paragraph()
    
    # Product
    doc.add_heading('Modelo: Product (Producto)', 3)
    doc.add_paragraph(
        'Representa un ítem del menú del restaurante. Relacionado con ingredientes y sabores.'
    )
    
    product_fields = [
        ('name', 'CharField(255)', 'Nombre del producto'),
        ('description', 'TextField', 'Descripción del producto'),
        ('price', 'DecimalField', 'Precio de venta'),
        ('cost_price', 'DecimalField', 'Precio de costo (calculado)'),
        ('category', 'CharField(20)', 'Categoría (promociones, entradas, principales, etc.)'),
        ('imagen', 'CharField(100)', 'URL de la imagen'),
        ('is_active', 'BooleanField', 'Si está disponible'),
        ('ingredientes', 'ManyToManyField', 'Relación con Ingredient'),
        ('sabores', 'ManyToManyField', 'Relación con Flavor')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Campo'
    hdr_cells[1].text = 'Tipo'
    hdr_cells[2].text = 'Descripción'
    
    for field, ftype, desc in product_fields:
        row_cells = table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = ftype
        row_cells[2].text = desc
    
    doc.add_paragraph()
    
    # Recipe (NUEVO)
    doc.add_heading('Modelo: Recipe (Receta) ⭐ NUEVO', 3)
    doc.add_paragraph(
        'Define la composición de un producto: qué ingredientes lleva y en qué cantidad. '
        'Permite el cálculo dinámico de costos cuando cambian los precios de ingredientes.'
    )
    
    recipe_fields = [
        ('product', 'ForeignKey(Product)', 'Producto al que pertenece'),
        ('ingredient', 'ForeignKey(Ingredient)', 'Ingrediente utilizado'),
        ('quantity', 'DecimalField', 'Cantidad necesaria del ingrediente')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Campo'
    hdr_cells[1].text = 'Tipo'
    hdr_cells[2].text = 'Descripción'
    
    for field, ftype, desc in recipe_fields:
        row_cells = table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = ftype
        row_cells[2].text = desc
    
    doc.add_paragraph()
    
    # InventoryMovement (NUEVO)
    doc.add_heading('Modelo: InventoryMovement (Movimiento de Inventario) ⭐ NUEVO', 3)
    doc.add_paragraph(
        'Registra todos los movimientos de inventario (compras, uso, ajustes, desperdicios). '
        'Permite análisis histórico con Pandas para proyección de compras futuras.'
    )
    
    movement_fields = [
        ('ingredient', 'ForeignKey(Ingredient)', 'Ingrediente afectado'),
        ('movement_type', 'CharField(20)', 'Tipo: PURCHASE, USAGE, ADJUSTMENT, WASTE'),
        ('quantity', 'DecimalField', 'Cantidad del movimiento'),
        ('cost_per_unit', 'DecimalField', 'Costo por unidad en este movimiento'),
        ('date', 'DateTimeField', 'Fecha del movimiento'),
        ('notes', 'TextField', 'Notas adicionales')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Campo'
    hdr_cells[1].text = 'Tipo'
    hdr_cells[2].text = 'Descripción'
    
    for field, ftype, desc in movement_fields:
        row_cells = table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = ftype
        row_cells[2].text = desc
    
    doc.add_paragraph()
    
    # Employee
    doc.add_heading('Modelo: Employee (Empleado)', 3)
    doc.add_paragraph('Gestión de recursos humanos del restaurante.')
    
    employee_fields = [
        ('name', 'CharField(100)', 'Nombre del empleado'),
        ('role', 'CharField(20)', 'Rol: CHEF, WAITER, DELIVERY, MANAGER, OTHER'),
        ('phone', 'CharField(20)', 'Teléfono de contacto'),
        ('salary_base', 'DecimalField', 'Salario base'),
        ('is_active', 'BooleanField', 'Si está activo'),
        ('joined_at', 'DateField', 'Fecha de ingreso')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Campo'
    hdr_cells[1].text = 'Tipo'
    hdr_cells[2].text = 'Descripción'
    
    for field, ftype, desc in employee_fields:
        row_cells = table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = ftype
        row_cells[2].text = desc
    
    doc.add_paragraph()
    doc.add_page_break()
    
    doc.add_heading('App: bot', 2)
    
    # Order
    doc.add_heading('Modelo: Order (Pedido)', 3)
    doc.add_paragraph('Representa una transacción de cliente desde Telegram o Web.')
    
    order_fields = [
        ('telegram_id', 'BigIntegerField', 'ID de Telegram del cliente'),
        ('item', 'CharField(255)', 'Nombre del producto pedido'),
        ('precio', 'DecimalField', 'Precio del pedido'),
        ('status', 'CharField(50)', 'Estado: pendiente, en_preparacion, listo, entregado'),
        ('service_type', 'CharField(20)', 'HERE (Comer Aquí) o TOGO (Para Llevar)'),
        ('delivery_mode', 'CharField(20)', 'PICKUP o DELIVERY'),
        ('location', 'TextField', 'Dirección de entrega'),
        ('payment_status', 'CharField(20)', 'Estado del pago'),
        ('payment_receipt_file_id', 'CharField(255)', 'ID del comprobante en Telegram'),
        ('fecha', 'DateTimeField', 'Fecha del pedido')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Campo'
    hdr_cells[1].text = 'Tipo'
    hdr_cells[2].text = 'Descripción'
    
    for field, ftype, desc in order_fields:
        row_cells = table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = ftype
        row_cells[2].text = desc
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # === 4. MÓDULOS DEL BACKEND ===
    doc.add_heading('4. Módulos del Backend', 1)
    
    modules = {
        'core': 'Gestión de productos, ingredientes, sabores, empleados y nómina. Incluye vistas de administración y estadísticas.',
        'bot': 'Lógica del Bot de Telegram, procesamiento de pedidos, webhooks y notificaciones.',
        'recomendacion': 'Motor de IA para recomendaciones personalizadas usando Machine Learning.',
        'ml': 'Utilidades de Machine Learning y procesamiento de lenguaje natural.',
        'analytics (NUEVO)': 'Automatizaciones con Pandas: recálculo de costos, matriz BCG, proyección de compras, reportes Excel.'
    }
    
    for module, description in modules.items():
        p = doc.add_paragraph()
        p.add_run(f'{module}: ').bold = True
        p.add_run(description)
    
    doc.add_page_break()
    
    # === 5. API REFERENCE ===
    doc.add_heading('5. API Reference Completo', 1)
    
    doc.add_heading('Endpoints de Productos', 2)
    
    product_endpoints = [
        ('GET', '/api/products/', 'Lista todos los productos agrupados por categoría'),
        ('GET', '/api/products-admin/', 'CRUD completo de productos (Admin)'),
        ('POST', '/api/products-admin/', 'Crear nuevo producto (Admin)'),
        ('PUT', '/api/products-admin/{id}/', 'Actualizar producto (Admin)'),
        ('DELETE', '/api/products-admin/{id}/', 'Eliminar producto (Admin)')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Método'
    hdr_cells[1].text = 'Endpoint'
    hdr_cells[2].text = 'Descripción'
    
    for method, endpoint, desc in product_endpoints:
        row_cells = table.add_row().cells
        row_cells[0].text = method
        row_cells[1].text = endpoint
        row_cells[2].text = desc
    
    doc.add_paragraph()
    
    doc.add_heading('Endpoints de Analytics (NUEVO) ⭐', 2)
    
    analytics_endpoints = [
        ('POST', '/api/analytics/recalculate-costs/', 'Recalcula costos de productos basándose en recetas'),
        ('GET', '/api/analytics/bcg-matrix/', 'Clasifica productos en Matriz BCG'),
        ('GET', '/api/analytics/purchase-suggestions/', 'Genera lista de compras sugeridas'),
        ('GET', '/api/analytics/export-excel/', 'Exporta reporte financiero completo en Excel')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Método'
    hdr_cells[1].text = 'Endpoint'
    hdr_cells[2].text = 'Descripción'
    
    for method, endpoint, desc in analytics_endpoints:
        row_cells = table.add_row().cells
        row_cells[0].text = method
        row_cells[1].text = endpoint
        row_cells[2].text = desc
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # === 6. AUTOMATIZACIONES CON PANDAS ===
    doc.add_heading('6. Automatizaciones con Pandas', 1)
    
    doc.add_paragraph(
        'El sistema incluye 4 automatizaciones clave implementadas con Pandas para '
        'optimizar las operaciones del restaurante:'
    )
    
    doc.add_heading('1. Control de Inventario Dinámico', 2)
    doc.add_paragraph(
        'Recalcula automáticamente el costo de todos los productos cuando cambian los '
        'precios de los ingredientes. Genera alertas de inflación y recomienda ajustes de precios.'
    )
    
    benefits_inventory = [
        'Detección instantánea de cambios en costos',
        'Alertas automáticas cuando la inflación supera el 5%',
        'Identificación de productos con margen bajo (<20%)',
        'Actualización masiva de precios de costo en milisegundos'
    ]
    
    for benefit in benefits_inventory:
        doc.add_paragraph(benefit, style='List Bullet')
    
    doc.add_heading('2. Matriz BCG (Ingeniería de Menú)', 2)
    doc.add_paragraph(
        'Clasifica automáticamente los productos en 4 categorías basándose en '
        'rentabilidad (margen) vs. popularidad (ventas):'
    )
    
    bcg_categories = {
        '⭐ Estrellas': 'Alta venta + Alto margen → Mantener y promocionar',
        '🐴 Caballos de Batalla': 'Alta venta + Bajo margen → Subir precio ligeramente',
        '❓ Incógnitas': 'Baja venta + Alto margen → Hacer publicidad',
        '🐶 Perros': 'Baja venta + Bajo margen → Considerar eliminar del menú'
    }
    
    for category, action in bcg_categories.items():
        p = doc.add_paragraph()
        p.add_run(f'{category}: ').bold = True
        p.add_run(action)
    
    doc.add_heading('3. Proyección de Compras', 2)
    doc.add_paragraph(
        'Analiza el histórico de ventas de los últimos 60 días y predice las necesidades '
        'de ingredientes para los próximos 7 días, considerando patrones por día de semana.'
    )
    
    benefits_prediction = [
        'Evita comprar de más (reduce desperdicio)',
        'Evita comprar de menos (previene pérdida de ventas)',
        'Optimiza el capital de trabajo',
        'Lista de compras automática con cantidades exactas y costos estimados'
    ]
    
    for benefit in benefits_prediction:
        doc.add_paragraph(benefit, style='List Bullet')
    
    doc.add_heading('4. Reportes Excel Avanzados', 2)
    doc.add_paragraph(
        'Exporta datos financieros en formato Excel con fórmulas nativas, '
        'formato profesional y múltiples hojas de cálculo.'
    )
    
    excel_sheets = [
        'Ventas Diarias: Con totales y promedios automáticos',
        'Nómina: Pagos a empleados con totales',
        'Productos BCG: Con cálculo de márgenes',
        'Ingredientes: Costos y unidades de medida'
    ]
    
    for sheet in excel_sheets:
        doc.add_paragraph(sheet, style='List Bullet')
    
    doc.add_page_break()
    
    # === 7. FRONTEND ===
    doc.add_heading('7. Frontend - Componentes y Páginas', 1)
    
    doc.add_heading('Páginas Principales', 2)
    
    pages = {
        'AdminPanel.jsx': 'Panel de administración completo con gestión de productos, ingredientes, empleados, estadísticas y analytics.',
        'KitchenPanel.jsx': 'Panel de cocina en tiempo real para gestión de pedidos activos.',
        'Login.jsx': 'Página de inicio de sesión con autenticación JWT y Telegram.',
        'Register.jsx': 'Registro de nuevos usuarios.',
        'Profile.jsx': 'Perfil de usuario con historial de pedidos.'
    }
    
    for page, desc in pages.items():
        p = doc.add_paragraph()
        p.add_run(f'{page}: ').bold = True
        p.add_run(desc)
    
    doc.add_heading('Componentes Clave', 2)
    
    components_list = [
        'MenuCategorias.jsx: Muestra productos organizados por categorías',
        'CarritoTelegram.jsx: Carrito de compras integrado con Telegram',
        'TelegramLogin.jsx: Widget de autenticación con Telegram',
        'PlatilloCard.jsx: Tarjeta de producto con imagen y detalles'
    ]
    
    for comp in components_list:
        doc.add_paragraph(comp, style='List Bullet')
    
    doc.add_page_break()
    
    # === 8. BOT DE TELEGRAM ===
    doc.add_heading('8. Bot de Telegram', 1)
    
    doc.add_paragraph(
        'El bot de Telegram permite a los clientes realizar pedidos de forma conversacional '
        'y recibir notificaciones en tiempo real.'
    )
    
    doc.add_heading('Comandos Disponibles', 2)
    
    commands = [
        '/start - Inicia la conversación con el bot',
        '/menu - Muestra el menú completo organizado por categorías',
        '/carrito - Muestra el carrito actual',
        '/soporte - Información de contacto para soporte',
        '/pedidos - Historial de pedidos del usuario'
    ]
    
    for cmd in commands:
        doc.add_paragraph(cmd, style='List Bullet')
    
    doc.add_heading('Flujo de Pedido', 2)
    
    flow_steps = [
        '1. Cliente solicita un producto (ej: "Quiero una hamburguesa")',
        '2. Bot muestra opciones disponibles con precios',
        '3. Cliente selecciona cantidad y sabor (si aplica)',
        '4. Bot solicita tipo de servicio (Comer Aquí / Para Llevar)',
        '5. Si es Para Llevar, solicita modo de entrega (Pickup / Delivery)',
        '6. Bot solicita comprobante de pago',
        '7. Admin verifica y aprueba el pago',
        '8. Pedido se envía a cocina',
        '9. Cliente recibe notificación cuando el pedido está listo'
    ]
    
    for step in flow_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_page_break()
    
    # === 9. MANUAL DE USUARIO ===
    doc.add_heading('9. Manual de Usuario', 1)
    
    doc.add_heading('Para Administradores', 2)
    
    doc.add_heading('Gestión de Productos', 3)
    admin_product_steps = [
        'Acceder a http://localhost:5173/admin-panel',
        'Iniciar sesión con credenciales de administrador',
        'En la sección "Productos", hacer clic en "Agregar Producto"',
        'Completar formulario: nombre, descripción, precio, costo, categoría, imagen',
        'Seleccionar ingredientes y sabores disponibles',
        'Guardar el producto'
    ]
    
    for step in admin_product_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_heading('Uso de Analytics (NUEVO)', 3)
    analytics_steps = [
        'En AdminPanel, navegar a la sección "Inteligencia de Negocio"',
        'Hacer clic en "Recalcular Costos" para actualizar precios de costo',
        'Revisar alertas de inflación y productos con margen bajo',
        'Consultar la Matriz BCG para ver clasificación de productos',
        'Revisar "Compras Sugeridas" antes de ir al mercado',
        'Hacer clic en "Exportar Excel" para descargar reporte completo'
    ]
    
    for step in analytics_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_heading('Para Personal de Cocina', 2)
    
    kitchen_steps = [
        'Acceder a http://localhost:5173/kitchen-panel',
        'Iniciar sesión con credenciales de cocina',
        'Ver pedidos activos organizados por tipo (Local, Delivery, Pickup)',
        'Marcar pedidos como "En Preparación" al comenzar',
        'Marcar pedidos como "Listo" al terminar',
        'El cliente recibirá notificación automática'
    ]
    
    for step in kitchen_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_page_break()
    
    # === 10. CONFIGURACIÓN Y DEPLOYMENT ===
    doc.add_heading('10. Configuración y Deployment', 1)
    
    doc.add_heading('Variables de Entorno', 2)
    doc.add_paragraph('Crear archivo .env en Backend/ con:')
    
    env_vars = [
        'SECRET_KEY=tu_clave_secreta_django',
        'DEBUG=True',
        'TELEGRAM_BOT_TOKEN=tu_token_de_telegram',
        'DATABASE_URL=sqlite:///db.sqlite3',
        'ALLOWED_HOSTS=localhost,127.0.0.1'
    ]
    
    for var in env_vars:
        doc.add_paragraph(var, style='List Bullet')
    
    doc.add_heading('Instalación Local', 2)
    
    install_steps = [
        'Clonar el repositorio',
        'Crear entorno virtual: python -m venv .venv',
        'Activar entorno: .venv\\Scripts\\activate (Windows)',
        'Instalar dependencias: pip install -r Backend/requirements.txt',
        'Ejecutar migraciones: python manage.py migrate',
        'Crear superusuario: python manage.py createsuperuser',
        'Iniciar servidor: python manage.py runserver',
        'En otra terminal, iniciar frontend: cd reactproject && npm install && npm run dev'
    ]
    
    for step in install_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_heading('Deployment en Producción', 2)
    
    prod_steps = [
        'Configurar PostgreSQL como base de datos',
        'Configurar variables de entorno en servidor',
        'Ejecutar collectstatic para archivos estáticos',
        'Configurar Nginx como reverse proxy',
        'Configurar Gunicorn para servir Django',
        'Configurar SSL con Let\'s Encrypt',
        'Configurar webhook de Telegram para producción'
    ]
    
    for step in prod_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_page_break()
    
    # === APÉNDICE ===
    doc.add_heading('Apéndice A: Troubleshooting', 1)
    
    issues = {
        'Error: ModuleNotFoundError': 'Asegurarse de que el entorno virtual está activado',
        'Error: CORS': 'Verificar configuración de CORS_ALLOWED_ORIGINS en settings.py',
        'Bot no responde': 'Verificar que el token de Telegram es correcto y el bot está iniciado',
        'Migraciones fallan': 'Eliminar db.sqlite3 y ejecutar migrate desde cero'
    }
    
    for issue, solution in issues.items():
        p = doc.add_paragraph()
        p.add_run(f'{issue}: ').bold = True
        p.add_run(solution)
    
    # === GUARDAR DOCUMENTO ===
    output_path = os.path.join(os.path.dirname(__file__), '..', 'Documentacion_Sistema_4Sabores.docx')
    doc.save(output_path)
    print(f'✅ Documentación generada exitosamente: {output_path}')
    return output_path

if __name__ == '__main__':
    create_documentation()
